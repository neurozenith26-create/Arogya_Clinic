"""Append HW/Infrastructure Assumptions appendix to the Arogya Clinic SoW.

Reads v1, appends Appendix A on a new page, writes v2. Idempotent — re-running
regenerates v2 from v1, so content edits live in this file.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "DZ2C_SOW_C001_Argya_Clinic-v1.docx"
DST = ROOT / "DZ2C_SOW_C001_Argya_Clinic-v4.docx"


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_rich_bullet(doc, parts):
    """parts = list of (text, bold) tuples — for inline bold."""
    p = doc.add_paragraph(style="List Bullet")
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
    return p


def build_appendix(doc):
    add_page_break(doc)
    add_heading(doc, "Appendix A — Hardware & Infrastructure Assumptions", level=1)

    add_para(
        doc,
        "This appendix records the HW/SaaS provisioning assumptions on which the "
        "Vendor's advisory work is predicated. Section 2 of this SoW places "
        "infrastructure provisioning out of scope; the items below describe what "
        "the Client is expected to procure and operate so that the architecture, "
        "design, and reviews delivered under this SoW can be realised. Detailed "
        "pricing is maintained separately in Arogya_Hosting_Proposal_QuickView.xlsx "
        "and is not duplicated here.",
        italic=True,
    )

    # A.1
    add_heading(doc, "A.1 Traffic & Load Baseline", level=2)
    add_para(doc, "The sizing on which this SoW's recommendations are based:")
    for b in [
        "300–800 daily unique patients · 200–300 peak concurrent users",
        "~20–60 GB / month frontend bandwidth",
        "5,000–15,000 daily database transactions",
        "Database volume grows to 5–15 GB over 3 years "
        "(includes BYTEA-stored reports & payment proofs)",
    ]:
        add_bullet(doc, b)
    add_para(
        doc,
        "Sized per Arogya_Hosting_Proposal_QuickView.xlsx, verified May 2026; "
        "to be re-baselined annually.",
        italic=True,
    )

    # A.2
    add_heading(doc, "A.2 Planned Production Stack — Option 3 Hybrid", level=2)
    add_para(
        doc,
        "The Client has selected Option 3 (Hybrid) per "
        "Arogya_Hosting_Proposal_QuickView.xlsx as the production stack to "
        "procure. Pre-production testing today runs on the FREE tiers of the "
        "same providers; this SoW's advisory work is sized for the paid "
        "production target below.",
    )
    stack_items = [
        ("Frontend", "Cloudflare Pages (free tier — unlimited bandwidth, unlimited requests, free DDoS + WAF, free auto-renewing SSL)"),
        ("Backend", "Hostinger KVM 2 Mumbai (≥ 2 vCPU / 8 GB RAM); Node.js 18+; Express on port 3001 behind the platform's reverse proxy"),
        ("Database", "Supabase Pro (Mumbai region) — managed Postgres 14+ with TLS enforced, daily automated backups, 7-day point-in-time recovery"),
        ("File / object storage", "Supabase Storage — private buckets reports, payment-proofs, invoices; public bucket doctor-photos. Backend reads/writes via the service-role key; the frontend never talks to Storage directly. Signed URLs are issued by the backend on demand for private files."),
        ("DNS / Domain", "Hostinger .com domain (free for Year 1 per the Hybrid option)"),
        ("SSL / TLS", "Cloudflare-managed at the edge + Let's Encrypt on the VPS (both auto-renew)"),
        ("Off-site DB backup", "Backblaze B2 / equivalent (~10 GB starter), encrypted at rest, retention ≥ 30 days"),
    ]
    for label, body in stack_items:
        add_rich_bullet(doc, [(f"{label} — ", True), (body, False)])

    # A.3
    add_heading(doc, "A.3 Payment Channel & Third-Party SaaS Accounts", level=2)
    add_para(
        doc,
        "Payments are accepted via a static UPI QR code displayed on the booking "
        "page. The patient pays directly from their UPI app, uploads the payment "
        "screenshot, and admin staff re-verify the receipt in person at the clinic "
        "(two-stage manual verification). No payment gateway, merchant account, or "
        "transaction processor is required for the current scope — patient money "
        "moves bank-to-bank and the Client incurs no per-transaction fee.",
    )
    add_para(doc, "Client must provide / procure:")
    for label, body in [
        ("UPI VPA + QR", "the clinic's UPI Virtual Payment Address (e.g. arogya@hdfcbank) and the corresponding QR image, configured in the admin payment settings"),
        ("Resend (email)", "account + verified sending domain with SPF, DKIM, DMARC records — used for ALL patient communications: OTP, booking-confirmed, report-ready, payment re-verified. Free tier covers 3,000 emails / month, comfortably above projected clinic volume."),
        ("Sentry", "project (optional; error reporting integration is supported)"),
    ]:
        add_rich_bullet(doc, [(f"{label} — ", True), (body, False)])
    add_para(
        doc,
        "SMS delivery (MSG91, Twilio, etc.) is explicitly NOT in scope — all "
        "patient notifications use email only. This keeps third-party costs at "
        "₹0 / month at current volume and avoids DLT registration overhead.",
        italic=True,
    )

    # A.4
    add_heading(doc, "A.4 Environments", level=2)
    add_para(
        doc,
        "Minimum two environments: Staging and Production. Each must have isolated "
        "database, isolated secrets, and isolated third-party API keys. Staging "
        "hosting cost is the Client's responsibility.",
    )

    # A.5
    add_heading(doc, "A.5 Production-Readiness Gates", level=2)
    add_para(
        doc,
        "The current MVP build carries one engineering choice that must be "
        "resolved before live patient traffic is opened. This is not a "
        "deferred-roadmap item — it is a pre-launch gate:",
        bold=True,
    )
    add_rich_bullet(doc, [
        ("OTP & session continuity — ", True),
        ("OTP delivery currently uses an in-memory map, so the backend can only run as a single instance. ", False),
        ("Before go-live, provision Redis (Upstash / Supabase Redis / equivalent) and switch to a Redis-backed OTP store", True),
        (" so the backend can be safely restarted or horizontally scaled without invalidating in-flight OTPs.", False),
    ])
    add_para(
        doc,
        "Status of the previously-flagged file-storage gate: RESOLVED. Patient "
        "reports, payment proofs, doctor photos, and invoices will be served "
        "from Supabase Storage on the production stack (Section A.2), eliminating "
        "the Postgres BYTEA growth risk. The application's pluggable "
        "StorageAdapter interface makes this swap a configuration change, not a "
        "code rewrite.",
        italic=True,
    )

    # A.6
    add_heading(doc, "A.6 Security, Backup, Compliance", level=2)
    add_para(doc, "Mandatory baselines that the Client must enforce in Production:")
    for b in [
        "HTTPS enforced end-to-end; HSTS enabled at host layer; HTTP redirects to HTTPS",
        "Database connection over TLS (DATABASE_SSL=true); credentials rotated on every team-member change",
        "JWT signing secret of at least 64 random characters; rotated on suspected compromise or staff exit",
        "bcrypt password-hashing cost factor ≥ 12 in Production (configurable via BCRYPT_ROUNDS)",
        "No patient PII (name, phone, report contents) written to application logs or third-party error trackers",
        "Daily automated database backups with 7-day point-in-time recovery retention minimum",
        "Off-site weekly DB dump to Backblaze B2 / equivalent, encrypted at rest, retention ≥ 30 days",
        "Secrets managed via Hostinger / Cloudflare / Supabase environment variables (no Vault/HSM tooling is included by Vendor)",
        "India-region hosting (Mumbai) for DPDP Act compliance on patient PII residency — both backend (Hostinger KVM Mumbai) and database (Supabase Mumbai)",
        "Cloudflare free WAF + unmetered DDoS protection accepted as baseline; rate-limiting enabled on login & OTP endpoints",
        "Production access (DB console, Hostinger panel, Cloudflare dashboard, Supabase Studio) restricted to ≤ 2 named operators with MFA enabled",
    ]:
        add_bullet(doc, b)

    # A.7
    add_heading(doc, "A.7 Operational Effort (Option 3 Hybrid)", level=2)
    add_para(
        doc,
        "Per Arogya_Hosting_Proposal_QuickView.xlsx, the Hybrid stack carries an "
        "ongoing operational load of approximately 3–4 hours per month for VPS "
        "administration on the Hostinger KVM backend — OS patching, Node.js "
        "version upgrades, log rotation, and similar housekeeping. This effort "
        "is the Client's responsibility unless contracted separately. Cloudflare "
        "Pages and Supabase Pro are fully managed and require zero operational "
        "effort.",
    )

    # A.8
    add_heading(doc, "A.8 Browser & Device Support", level=2)
    add_para(
        doc,
        "Evergreen browsers (Chrome, Edge, Safari, Firefox — last two major versions). "
        "Responsive layout supported down to 360 px viewport width. Legacy Internet "
        "Explorer and Android < 8 are explicitly not supported.",
    )

    # A.9
    add_heading(doc, "A.9 Explicit Exclusions", level=2)
    add_para(doc, "Reinforcing Section 2 — the following are not covered:")
    for b in [
        "Payment-gateway integration (Razorpay / Stripe / PayU / Cashfree) — current scope is manual UPI QR only; gateway integration is a future change-request item",
        "CI/CD pipeline setup beyond the platform defaults of Cloudflare Pages / Render",
        "Vault / HSM / SIEM / dedicated SOC",
        "Advanced WAF rule authoring beyond the Cloudflare free-tier 5-rule allowance",
        "24×7 on-call operations, performance load testing, penetration testing",
        "Migration of historical patient data from any legacy system",
    ]:
        add_bullet(doc, b)

    add_para(
        doc,
        "Pricing for each line item above is in Arogya_Hosting_Proposal_QuickView.xlsx "
        "(Cost Details sheet). This SoW does not quote infrastructure figures; they "
        "are reviewed and approved separately.",
        italic=True,
    )


def main():
    if not SRC.exists():
        raise SystemExit(f"Source SoW not found: {SRC}")
    doc = Document(str(SRC))
    build_appendix(doc)
    doc.save(str(DST))
    print(f"Written: {DST}")


if __name__ == "__main__":
    main()
