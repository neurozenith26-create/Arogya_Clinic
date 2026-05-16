"""Arogya Diagnostics — CLIENT-READY hosting proposal (3 approaches).
Sized for GOING LIVE: 1,000-2,000 unique users/day + up to 1,000 concurrent peaks.
All prices verified May 2026 via direct web search."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Colors ----------
TEAL = RGBColor(0x0F, 0x76, 0x6E)
TEAL_DARK = RGBColor(0x0B, 0x4F, 0x4A)
ORANGE = RGBColor(0xF9, 0x73, 0x16)
GOLD = RGBColor(0xF5, 0x9E, 0x0B)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED = RGBColor(0xDC, 0x26, 0x26)
SLATE = RGBColor(0x33, 0x41, 0x55)
GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FE = RGBColor(0x0E, 0xA5, 0xE9)
BE = RGBColor(0x8B, 0x5C, 0xF6)
DB = RGBColor(0x05, 0x96, 0x69)

HEX_TEAL = "0F766E"
HEX_TEAL_L = "CCFBF1"
HEX_ORANGE = "F97316"
HEX_ORANGE_L = "FFEDD5"
HEX_GOLD = "F59E0B"
HEX_GOLD_L = "FEF3C7"
HEX_GRAY_L = "F1F5F9"
HEX_FE = "0EA5E9"
HEX_FE_L = "E0F2FE"
HEX_BE = "8B5CF6"
HEX_BE_L = "EDE9FE"
HEX_DB = "059669"
HEX_DB_L = "D1FAE5"
HEX_GREEN_L = "DCFCE7"
HEX_RED_L = "FEE2E2"


def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color)
    tc_pr.append(s)


def borders(cell, color="0F766E", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), size)
        e.set(qn('w:color'), color)
        b.append(e)
    tc_pr.append(b)


def shade_para(p, hex_color):
    p_pr = p._p.get_or_add_pPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color)
    p_pr.append(s)


def run(p, text, *, bold=False, size=11, color=SLATE, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return r


def banner(doc, text, *, fill=HEX_TEAL, size=18, color=WHITE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_para(p, fill)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run(p, "  " + text + "  ", bold=True, size=size, color=color)


def heading(doc, text, *, color=TEAL, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run(p, text, bold=True, size=size, color=color)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '0F766E')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def sub_heading(doc, text, *, color=ORANGE, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run(p, text, bold=True, size=size, color=color)


def body(doc, text, *, size=11, color=SLATE, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run(p, text, bold=bold, size=size, color=color, italic=italic)


def bullet(doc, text, *, size=11, marker_color=TEAL, color=SLATE, marker="•"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    run(p, marker + "  ", bold=True, color=marker_color, size=size)
    run(p, text, size=size, color=color)


def info_box(doc, label, body_text, *, label_color=GOLD, fill=HEX_GOLD_L,
             border="F59E0B"):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, fill)
    borders(cell, color=border, size="10")
    cell.text = ""
    p = cell.paragraphs[0]
    run(p, label, bold=True, size=12, color=label_color)
    p2 = cell.add_paragraph()
    run(p2, body_text, size=11, color=SLATE)
    doc.add_paragraph()


def architecture_table(doc, rows, *, accent_hex=HEX_TEAL):
    """rows: list of (tag_label, tag_hex_color, fill_hex, where, what, cost)"""
    tbl = doc.add_table(rows=len(rows) + 1, cols=4)
    hdrs = ["Layer", "Hosted on", "What it does", "Verified cost"]
    for i, h in enumerate(hdrs):
        cell = tbl.rows[0].cells[i]
        shade(cell, accent_hex)
        cell.text = ""
        run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
    for r_idx, (tag, c_dark, c_light, where, what, cost) in enumerate(rows, start=1):
        cell0 = tbl.rows[r_idx].cells[0]
        shade(cell0, c_light)
        cell0.text = ""
        run(cell0.paragraphs[0], tag, bold=True, size=10,
            color=RGBColor.from_string(c_dark))
        cell1 = tbl.rows[r_idx].cells[1]
        cell1.text = ""
        run(cell1.paragraphs[0], where, bold=True, size=10, color=SLATE)
        cell2 = tbl.rows[r_idx].cells[2]
        cell2.text = ""
        run(cell2.paragraphs[0], what, size=9, color=SLATE)
        cell3 = tbl.rows[r_idx].cells[3]
        cell3.text = ""
        run(cell3.paragraphs[0], cost, size=9, color=TEAL, bold=True)
    doc.add_paragraph()


def pricing_table(doc, rows, totals_row, *, total_fill=HEX_GOLD_L,
                  total_color=GOLD):
    """rows: list of (tag, tag_hex, item_text, y1, y2, y3)
       totals_row: (label, y1_total, y2_total, y3_total)
    """
    tbl = doc.add_table(rows=len(rows) + 1, cols=4)
    hdrs = ["Item", "Year 1", "Year 2", "Year 3"]
    for i, h in enumerate(hdrs):
        cell = tbl.rows[0].cells[i]
        shade(cell, HEX_TEAL)
        cell.text = ""
        run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
    for r_idx, (tag, color_hex, item, y1, y2, y3) in enumerate(rows, start=1):
        cell0 = tbl.rows[r_idx].cells[0]
        cell0.text = ""
        p = cell0.paragraphs[0]
        run(p, "[" + tag + "] ", bold=True, size=9,
            color=RGBColor.from_string(color_hex))
        run(p, item, size=10, color=SLATE)
        for c_idx, val in enumerate([y1, y2, y3], start=1):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = ""
            run(cell.paragraphs[0], val, size=10, color=SLATE)
            if r_idx % 2 == 1:
                shade(cell, HEX_GRAY_L)
        if r_idx % 2 == 1:
            shade(cell0, HEX_GRAY_L)

    doc.add_paragraph()

    tbl_t = doc.add_table(rows=2, cols=4)
    for i, h in enumerate(["", "Year 1", "Year 2", "Year 3"]):
        cell = tbl_t.rows[0].cells[i]
        shade(cell, HEX_GOLD)
        cell.text = ""
        run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
    for i, val in enumerate(totals_row):
        cell = tbl_t.rows[1].cells[i]
        shade(cell, total_fill)
        cell.text = ""
        run(cell.paragraphs[0], val, bold=True, size=12, color=total_color)
    doc.add_paragraph()


def pros_cons(doc, pros, cons):
    pc = doc.add_table(rows=1, cols=2)
    pcell = pc.rows[0].cells[0]
    ccell = pc.rows[0].cells[1]
    shade(pcell, HEX_GREEN_L)
    shade(ccell, HEX_RED_L)
    borders(pcell, color="16A34A", size="8")
    borders(ccell, color="DC2626", size="8")
    pcell.text = ""
    run(pcell.paragraphs[0], "PROS", bold=True, size=12, color=GREEN)
    for p_text in pros:
        pp = pcell.add_paragraph()
        pp.paragraph_format.left_indent = Cm(0.3)
        run(pp, "+  ", bold=True, color=GREEN, size=11)
        run(pp, p_text, size=10, color=SLATE)
    ccell.text = ""
    run(ccell.paragraphs[0], "CONS", bold=True, size=12, color=RED)
    for c_text in cons:
        pp = ccell.add_paragraph()
        pp.paragraph_format.left_indent = Cm(0.3)
        run(pp, "-  ", bold=True, color=RED, size=11)
        run(pp, c_text, size=10, color=SLATE)
    doc.add_paragraph()


# ===================================================================
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.4)
    s.bottom_margin = Cm(1.4)
    s.left_margin = Cm(1.8)
    s.right_margin = Cm(1.8)

# ===================================================================
# COVER PAGE
# ===================================================================
banner(doc, "AROGYA DIAGNOSTICS  ·  Hosting & Infrastructure Proposal",
       fill=HEX_TEAL, size=17)
banner(doc, "GOING-LIVE Plan  ·  Sized for 1,000 concurrent users  ·  3-Year Costs  ·  Verified May 2026",
       fill=HEX_ORANGE, size=11)
doc.add_paragraph()

info_box(doc,
    "WHAT THIS DOCUMENT IS",
    "Four hosting options for Arogya's website, custom APIs, and patient "
    "database over the next 3 years — sized for GOING LIVE with up to "
    "1,000 patients using the site at the same time (peak concurrency) "
    "and 1,000-2,000 unique users per day. Approach 0 is the current "
    "free-tier stack (Rs 0/year — will NOT survive at this scale, dev "
    "baseline only). Approaches 1, 2, 3 are real production-grade choices "
    "with comfortable headroom for 1,000 concurrent. Every rupee figure "
    "was verified directly from the provider's official website in May 2026.")

# Scale snapshot
heading(doc, "1.  What we're sizing for  ·  GOING-LIVE scale")
body(doc, "Hosting costs scale with traffic. These numbers reflect the going-live target — "
          "no soft-launch — so every approach below has comfortable headroom for the peaks:")
tbl = doc.add_table(rows=5, cols=2)
specs = [
    ("Daily unique patients", "1,000 - 2,000 per day"),
    ("Peak concurrent users (same time)", "up to 1,000 people"),
    ("Monthly bandwidth (Frontend)", "~60-150 GB/month (static React site)"),
    ("Daily database transactions", "~20,000-50,000 inserts/updates"),
    ("Database size after 3 years (with reports)", "~10-30 GB"),
]
for r_idx, (k, v) in enumerate(specs):
    cell_k = tbl.rows[r_idx].cells[0]
    cell_v = tbl.rows[r_idx].cells[1]
    shade(cell_k, HEX_TEAL_L)
    cell_k.text = ""
    run(cell_k.paragraphs[0], k, bold=True, size=10, color=TEAL)
    cell_v.text = ""
    run(cell_v.paragraphs[0], v, size=10, color=SLATE)
doc.add_paragraph()


# ===================================================================
# WHAT IS WHAT
# ===================================================================
heading(doc, "2.  The 3 pieces of the application — simple words")
body(doc,
     "Think of the application like a restaurant. Each piece needs its own "
     "place on the internet.")
parts = [
    ("[FRONTEND]", "The dining room — what patients SEE",
     "React website: homepage, booking forms, dashboard. Static files only.",
     HEX_FE, HEX_FE_L, FE),
    ("[BACKEND]", "The kitchen — our CUSTOM APIs",
     "Node.js + Express code WE built. Handles every login, booking, "
     "payment, report upload. No third-party API service — we own the code.",
     HEX_BE, HEX_BE_L, BE),
    ("[DATABASE]", "The pantry/safe — all patient data",
     "PostgreSQL stores names, mobiles, bookings, medical reports. The most "
     "sensitive layer. DPDP Act 2023 makes the clinic legally responsible "
     "for protecting this.",
     HEX_DB, HEX_DB_L, DB),
]
for tag, head, text, dark, light, color in parts:
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, light)
    borders(cell, color=dark, size="12")
    cell.text = ""
    p = cell.paragraphs[0]
    run(p, tag + "  ", bold=True, size=13, color=color)
    run(p, head, bold=True, size=11, color=SLATE)
    p2 = cell.add_paragraph()
    run(p2, text, size=10, color=SLATE)
    doc.add_paragraph()

doc.add_page_break()


# ===================================================================
# OVERVIEW — Current + 3 Approaches (side-by-side preview)
# ===================================================================
banner(doc, "AT-A-GLANCE  ·  Current Setup vs 3 Approaches",
       fill=HEX_ORANGE, size=15)
doc.add_paragraph()

tbl = doc.add_table(rows=5, cols=5)
hdrs = ["",
        "Approach 0\n(Current Stack)",
        "Approach 1\n(Premium Cloud)",
        "Approach 2\n(Single VPS)",
        "Approach 3 ★\n(Hybrid — Recommended)"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, h, bold=True, size=9, color=WHITE)

overview_rows = [
    ("[FRONTEND]", HEX_FE, HEX_FE_L,
     "Netlify Free\n(.netlify.app)",
     "Cloudflare Pages\n(Free, unlimited BW)",
     "On the same Hostinger\nKVM 4 VPS",
     "Cloudflare Pages\n(Free, unlimited BW)"),
    ("[BACKEND]", HEX_BE, HEX_BE_L,
     "Render Free\n(sleeps 15 min)",
     "Render Pro\n$85/mo (2 CPU/4 GB)",
     "On the same Hostinger\nKVM 4 VPS (Mumbai)",
     "Hostinger KVM 4 VPS\n(Mumbai, 4 CPU/16 GB)"),
    ("[DATABASE]", HEX_DB, HEX_DB_L,
     "Supabase Free\n(Mumbai, pauses!)",
     "Supabase Pro\n$25/mo (Mumbai)",
     "PostgreSQL on SAME\nKVM 4 (Mumbai)",
     "Supabase Pro\n$25/mo (Mumbai)"),
    ("3-YEAR TOTAL", HEX_GOLD, HEX_GOLD_L,
     "Rs 0\n(will CRASH live)",
     "Rs 3,41,233",
     "Rs 69,524",
     "Rs 1,46,028  ★"),
]
for r_idx, (label, c_dark, c_light, *vals) in enumerate(overview_rows,
                                                         start=1):
    cell = tbl.rows[r_idx].cells[0]
    shade(cell, c_light)
    cell.text = ""
    run(cell.paragraphs[0], label, bold=True, size=9,
        color=RGBColor.from_string(c_dark))
    for c_idx, val in enumerate(vals, start=1):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c_idx == 4:
            shade(cell, HEX_GOLD_L)
            run(p, val, bold=(r_idx == 4), size=9,
                color=GOLD if r_idx == 4 else SLATE)
        else:
            run(p, val, bold=(r_idx == 4), size=9,
                color=TEAL if r_idx == 4 else SLATE)

doc.add_paragraph()

info_box(doc,
    "QUICK READ  ·  GOING LIVE FOR 1,000 CONCURRENT",
    "Approach 0 is what is running today — Rs 0 cost, but the Render Free "
    "backend will crash within seconds of going live (cannot serve 1,000 "
    "concurrent). Approach 1 is fully managed but most expensive — Render "
    "Pro at $85/mo handles 1,000 concurrent comfortably. Approach 2 is "
    "cheapest viable production option but pushes a single KVM 4 box to "
    "its limit at 1,000 concurrent and co-locates everything on one server "
    "(risky for patient data). Approach 3 is the balanced winner: KVM 4 "
    "Mumbai for the backend (4 vCPU / 16 GB — comfortable headroom for "
    "1,000 concurrent on Node alone since DB is offloaded) + Supabase Pro "
    "Mumbai for the DB. Saves Rs 1,95,000 vs Approach 1 at this scale.",
    label_color=ORANGE, fill=HEX_ORANGE_L, border="F97316")

doc.add_page_break()


# ===================================================================
# APPROACH 0 — Current Stack (what's actually running today)
# ===================================================================
banner(doc, "APPROACH 0  ·  Current Stack (free-tier — what's running today)",
       fill=HEX_TEAL, size=16)
doc.add_paragraph()

sub_heading(doc,
    "Stack: Netlify Free (Frontend) + Render Free (Backend) + Supabase Free Mumbai (Database)",
    color=TEAL)

body(doc,
     "This is exactly what is deployed at this moment — every layer on a "
     "100% free plan, no money spent. Frontend lives at "
     "arogya-clinic.netlify.app, backend at arogya-clinic.onrender.com, "
     "and the Postgres database is on Supabase's Mumbai (ap-south-1) "
     "region but on the free plan. Total monthly cost today: Rs 0. "
     "Fine for development and demos, but going LIVE with 1,000 concurrent "
     "users on this stack would crash within minutes — the cons below "
     "explain exactly why.")

info_box(doc,
    "⚠  WILL CRASH AT GOING-LIVE SCALE",
    "Render Free (0.5 CPU + 512 MB RAM, sleeps after 15 min idle) cannot "
    "serve 1,000 concurrent users — it caps out at ~30-50 concurrent in "
    "practice. Supabase Free has a hard ~60 direct / ~200 pooled "
    "connection limit and auto-pauses after 7 days. Netlify Free's 100 GB "
    "bandwidth ceiling would be exceeded mid-month at this scale. The "
    "good news: data is already in Mumbai (Supabase ap-south-1), so the "
    "migration to Approach 1 or 3 is just an in-place plan upgrade — no "
    "data export/import. Architecture stays the same; only the plan tier "
    "changes.",
    label_color=RED, fill=HEX_RED_L, border="DC2626")

heading(doc, "Where each piece lives", color=TEAL)
architecture_table(doc, [
    ("[FRONTEND]", HEX_FE, HEX_FE_L,
     "Netlify Free",
     "React build on Netlify global CDN at arogya-clinic.netlify.app. 100 GB bandwidth + 300 build min/month.",
     "Rs 0 / year"),
    ("[BACKEND]", HEX_BE, HEX_BE_L,
     "Render Free",
     "Node API at arogya-clinic.onrender.com. SLEEPS after 15 min idle — 30-50 sec cold start on first request.",
     "Rs 0 / year"),
    ("[DATABASE]", HEX_DB, HEX_DB_L,
     "Supabase Free (Mumbai)",
     "Postgres in AWS ap-south-1. 500 MB DB, 1 GB egress. AUTO-PAUSES after 7 days inactivity. No daily backup, no PITR.",
     "Rs 0 / year"),
    ("[DOMAIN]", "92400E", HEX_GOLD_L,
     "None (using subdomains)",
     "Using free .netlify.app and .onrender.com subdomains. No custom domain purchased yet.",
     "Rs 0 / year"),
])

heading(doc, "3-Year cost breakdown", color=TEAL)
pricing_table(doc, [
    ("FE", HEX_FE,
     "Frontend — Netlify Free (100 GB bandwidth limit — TIGHT for our scale)",
     "Rs 0", "Rs 0", "Rs 0"),
    ("BE", HEX_BE,
     "Backend — Render Free (sleeps 15 min idle; 30-50 sec cold start)",
     "Rs 0", "Rs 0", "Rs 0"),
    ("DB", HEX_DB,
     "Database — Supabase Free Mumbai (500 MB, auto-pauses, no PITR)",
     "Rs 0", "Rs 0", "Rs 0"),
    ("DOMAIN", "92400E",
     "Domain — none yet (using free .netlify.app + .onrender.com subdomains)",
     "Rs 0", "Rs 0", "Rs 0"),
    ("EXTRA", "64748B",
     "Off-site backups — none configured",
     "Rs 0", "Rs 0", "Rs 0"),
    ("EXTRA", "64748B",
     "DDoS / WAF — none (relying on platform defaults)",
     "Rs 0", "Rs 0", "Rs 0"),
], ["TOTAL per year", "Rs 0", "Rs 0", "Rs 0"])

info_box(doc,
    "3-YEAR TOTAL  ·  Rs 0  (development-only baseline)",
    "Zero cost — but also zero production guarantees. Suitable for "
    "demos, internal testing, and the first ~20-50 daily users. Before "
    "real patients touch the system, upgrade to Approach 1 or Approach 3. "
    "Because Supabase is already in Mumbai (ap-south-1), the Approach 3 "
    "migration is a same-region plan upgrade (Free → Pro) — no data move.",
    label_color=TEAL, fill=HEX_TEAL_L, border="0F766E")

heading(doc, "Pros & Cons", color=TEAL)
pros_cons(doc, pros=[
    "Zero rupees spent — completely free across all three layers.",
    "Already deployed and live (arogya-clinic.netlify.app + .onrender.com).",
    "Supabase Free is already in Mumbai (ap-south-1) — same region as Approach 1/3.",
    "Great for development, demos, and the first ~20-50 daily users.",
    "Migration path is in-place upgrades (Render Free -> Standard, Supabase Free -> Pro).",
    "Backend, frontend, DB are already separated — same architecture as Approach 1/3.",
], cons=[
    "Render Free SLEEPS after 15 min idle — first patient of the day waits 30-50 sec for cold start.",
    "Supabase Free AUTO-PAUSES after 7 days inactivity — site can go fully offline.",
    "Supabase Free has NO daily backups and NO PITR — a delete-by-accident is permanent.",
    "Supabase Free is only 500 MB DB + 1 GB egress — too small for stored report PDFs.",
    "Netlify Free 100 GB bandwidth/month is too close to our ~90 GB projection — risk of being paused mid-month.",
    "No DDoS protection, no WAF — direct exposure to script-kiddie attacks.",
    "No compliance certifications visible — can't show DPDP audit readiness.",
    "No custom domain — patients see arogya-clinic.netlify.app, not arogyadiagnostics.com.",
])

doc.add_page_break()


# ===================================================================
# APPROACH 1 — Premium All-Managed Cloud
# ===================================================================
banner(doc, "APPROACH 1  ·  Premium All-Managed Cloud",
       fill=HEX_TEAL, size=16)
doc.add_paragraph()

sub_heading(doc,
    "Stack: Cloudflare Pages + Render Pro + Supabase Pro Mumbai + Hostinger Domain",
    color=TEAL)

body(doc,
     "Three specialist managed platforms — each tier handed off to a "
     "service that handles scaling, backups, security patches, and uptime. "
     "Zero server administration: every deploy is git push, scaling is "
     "automatic, patches happen behind the scenes. Best developer "
     "experience, highest monthly cost. Sized with Render Pro (2 CPU / "
     "4 GB RAM) so the backend comfortably absorbs 1,000 concurrent users "
     "in Node cluster mode.")

info_box(doc,
    "WHY RENDER PRO ($85/mo) AND NOT STANDARD ($25/mo)",
    "Render Standard is 1 CPU + 2 GB RAM — fine up to ~300-500 concurrent, "
    "but it WILL crack at 1,000 concurrent (Node single-process saturates a "
    "single core quickly under high I/O fan-out). Render Pro is 2 CPU + "
    "4 GB RAM, which lets PM2 cluster across both cores and easily handles "
    "1,000 concurrent on the API workload. Pro Plus ($185/mo, 4 CPU / 8 GB) "
    "is one click away if traffic doubles.",
    label_color=ORANGE, fill=HEX_ORANGE_L, border="F97316")

heading(doc, "Where each piece lives", color=TEAL)
architecture_table(doc, [
    ("[FRONTEND]", HEX_FE, HEX_FE_L,
     "Cloudflare Pages",
     "Hosts React build on global CDN. Mumbai edge serves Indian patients. Unlimited bandwidth.",
     "Rs 0 (free, unlimited)"),
    ("[BACKEND]", HEX_BE, HEX_BE_L,
     "Render Pro ($85/mo)",
     "2 CPU + 4 GB RAM dedicated container. PM2 cluster across 2 cores. Handles 1,000 concurrent comfortably.",
     "Rs 86,700/year"),
    ("[DATABASE]", HEX_DB, HEX_DB_L,
     "Supabase Pro Mumbai ($25/mo)",
     "Managed Postgres in AWS Mumbai. 8 GB DB, 100 GB storage, daily backup, 7-day PITR, PgBouncer pooler handles 1,000 concurrent.",
     "Rs 25,500/year"),
    ("[DOMAIN]", "92400E", HEX_GOLD_L,
     "Hostinger registrar",
     ".com domain pointed via Cloudflare DNS. ",
     "Rs 199 Y1, Rs 1,299 renewal"),
])

heading(doc, "3-Year cost breakdown", color=TEAL)
pricing_table(doc, [
    ("FE", HEX_FE, "Frontend — Cloudflare Pages (free, unlimited)",
     "Rs 0", "Rs 0", "Rs 0"),
    ("BE", HEX_BE, "Backend — Render Pro $85/mo (2 CPU / 4 GB, handles 1,000 concurrent)",
     "Rs 86,700", "Rs 86,700", "Rs 86,700"),
    ("DB", HEX_DB, "Database — Supabase Pro Mumbai $25/mo",
     "Rs 25,500", "Rs 25,500", "Rs 25,500"),
    ("DOMAIN", "92400E", "Domain — .com (Rs 199 promo Y1)",
     "Rs 199", "Rs 1,299", "Rs 1,299"),
    ("EXTRA", "64748B", "Off-site backup — Backblaze B2 ~10 GB",
     "Rs 612", "Rs 612", "Rs 612"),
    ("EXTRA", "64748B", "Email — Resend free tier (3,000/mo free)",
     "Rs 0", "Rs 0", "Rs 0"),
], ["TOTAL per year", "Rs 1,13,011", "Rs 1,14,111", "Rs 1,14,111"])

info_box(doc,
    "3-YEAR TOTAL  ·  Rs 3,41,233  (~Rs 9,479/month average)",
    "Most expensive but the least operational work — there is no server to "
    "patch, no PM2 to monitor, no backup cron to maintain. Suited if Arogya "
    "has zero in-house tech team and wants every layer professionally managed. "
    "At 1,000 concurrent, Render Pro is the right tier — Standard would "
    "crack and Pro Plus is overkill until traffic doubles.",
    label_color=TEAL, fill=HEX_TEAL_L, border="0F766E")

heading(doc, "Pros & Cons", color=TEAL)
pros_cons(doc, pros=[
    "Fully managed — zero devops; everything auto-scales and auto-patches.",
    "Render Pro (2 CPU / 4 GB) comfortably handles 1,000 concurrent in Node cluster mode.",
    "Best uptime (~99.95%); each tier is a multi-AZ managed service.",
    "Built-in audit logs, monitoring, alerting on every layer.",
    "Render auto-deploys on git push — CI/CD is effectively free.",
    "Easy to hand off — any developer in India can pick this up.",
    "One-click upgrade path to Pro Plus ($185/mo, 4 CPU / 8 GB) if traffic doubles.",
], cons=[
    "Most expensive: ~Rs 3.4 lakh over 3 years — Rs 1.95 lakh more than Approach 3.",
    "Render Pro at $85/mo is ~6.5x the cost of an equivalent Hostinger KVM 4 in Mumbai.",
    "Render bills are in USD — small FX risk to the client.",
    "Three separate vendor accounts and three separate invoices.",
    "Egress (bandwidth out of Supabase) is metered and can spike.",
    "Render backend is hosted in US/EU regions — only the DB is in Mumbai.",
])

doc.add_page_break()


# ===================================================================
# APPROACH 2 — Single Hostinger KVM 4 VPS (All-in-one)
# ===================================================================
banner(doc, "APPROACH 2  ·  Single Hostinger KVM 4 VPS (All-in-one)",
       fill=HEX_TEAL, size=16)
doc.add_paragraph()

sub_heading(doc,
    "Stack: One Hostinger KVM 4 VPS (Mumbai) runs Frontend + Backend + Postgres + Nginx",
    color=TEAL)

body(doc,
     "ONE virtual server in Mumbai does everything: Nginx serves the React "
     "build as static files, PM2 runs the Node backend on port 3001, and "
     "PostgreSQL runs on localhost. One vendor, one invoice in INR, one "
     "SSH login. Cheapest of the three approaches. The trade-off: if that "
     "ONE box is compromised, the attacker has the Frontend AND Backend "
     "AND Database all at once.")

info_box(doc,
    "WHY KVM 4 — AND WHY IT'S TIGHT AT 1,000 CONCURRENT",
    "KVM 4 = 4 vCPU + 16 GB RAM + 200 GB NVMe + 16 TB bandwidth. At 1,000 "
    "concurrent users the box must run Node (cluster mode using ~2 cores) "
    "AND Postgres (4-8 GB RAM for working set, ~2 cores under load) AND "
    "Nginx — all on the same 4-core machine. This works but with little "
    "headroom; expect to monitor closely and upgrade to KVM 8 "
    "(~Rs 2,199/mo promo · Rs 4,399 renewal · 8 vCPU / 32 GB) if CPU "
    "sustains >70%. For 1,000 concurrent with comfortable headroom, "
    "Approach 3 (KVM 4 backend + offloaded Supabase DB) is the safer pick.",
    label_color=ORANGE, fill=HEX_ORANGE_L, border="F97316")

heading(doc, "Where each piece lives", color=TEAL)
architecture_table(doc, [
    ("[FRONTEND]", HEX_FE, HEX_FE_L,
     "Same KVM 4 VPS",
     "Nginx serves React build files from /var/www/arogya",
     "Bundled in VPS price"),
    ("[BACKEND]", HEX_BE, HEX_BE_L,
     "Same KVM 4 VPS",
     "Node APIs via PM2 cluster on port 3001, reverse-proxied by Nginx",
     "Bundled in VPS price"),
    ("[DATABASE]", HEX_DB, HEX_DB_L,
     "Same KVM 4 VPS",
     "PostgreSQL on localhost:5432, NOT exposed to public internet",
     "Bundled in VPS price"),
    ("[DOMAIN]", "92400E", HEX_GOLD_L,
     "Hostinger registrar",
     "Free .com first year (bundled with VPS plan)",
     "Rs 0 Y1, Rs 1,299 renewal"),
])

heading(doc, "3-Year cost breakdown", color=TEAL)
pricing_table(doc, [
    ("VPS", HEX_BE,
     "Hostinger KVM 4 (4 vCPU, 16 GB RAM, Mumbai DC) — 24-mo plan, incl. 18% GST",
     "Rs 15,562", "Rs 15,562", "Rs 33,966"),
    ("DOMAIN", "92400E",
     "Domain — .com (free Y1 with VPS, then renewal)",
     "Rs 0", "Rs 1,299", "Rs 1,299"),
    ("SEC", "DC2626",
     "SSL — Let's Encrypt (free, auto-renew)",
     "Rs 0", "Rs 0", "Rs 0"),
    ("EXTRA", "64748B",
     "Off-site DB backup — Backblaze B2 ~10 GB",
     "Rs 612", "Rs 612", "Rs 612"),
    ("EXTRA", "64748B",
     "Email — Resend free tier",
     "Rs 0", "Rs 0", "Rs 0"),
], ["TOTAL per year", "Rs 16,174", "Rs 17,473", "Rs 35,877"])

info_box(doc,
    "3-YEAR TOTAL  ·  Rs 69,524  (~Rs 1,931/month average)",
    "By far the cheapest option. But the savings come with risks — see "
    "the cons below. Suitable only if the clinic has someone (in-house or "
    "outsourced) who can do basic Linux administration, OR is willing to "
    "accept the security trade-off of co-locating all 3 layers.",
    label_color=TEAL, fill=HEX_TEAL_L, border="0F766E")

heading(doc, "Pros & Cons", color=TEAL)
pros_cons(doc, pros=[
    "Cheapest: Rs 69,524 over 3 years — saves Rs 76,000+ vs Approach 3.",
    "Indian pricing in INR, paid by UPI / NetBanking — no FX surprises.",
    "ONE vendor, ONE invoice — easy for a non-technical owner.",
    "Full root access — install anything (Redis, Puppeteer for PDF reports, etc).",
    "Hostinger Mumbai DC keeps all data in India (DPDP residency).",
    "Vertically scalable: KVM 4 -> KVM 8 (~Rs 4,399/mo renewal) on the same plan.",
], cons=[
    "TIGHT at 1,000 concurrent — Node + Postgres + Nginx all share 4 vCPUs; needs careful tuning.",
    "Likely needs upgrade to KVM 8 within 6-12 months if usage actually hits the 1,000 peak (~Rs 2,199-4,399/mo).",
    "ONE box runs everything — if hacked, ENTIRE patient database is leaked.",
    "Operating-system patching is the owner's responsibility — risky if missed.",
    "Backups are manual: pg_dump cron + Backblaze sync must be set up correctly.",
    "Renewal pricing more than doubles (Rs 1,099 -> Rs 2,399/mo) in Year 3.",
    "No managed compliance certifications — no SOC 2 / ISO 27001 report we can show auditors.",
    "Single point of failure: if the VPS dies, EVERYTHING goes down until restore.",
])

doc.add_page_break()


# ===================================================================
# APPROACH 3 — Hybrid (Recommended)
# ===================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_para(p, HEX_GOLD)
p.paragraph_format.space_after = Pt(2)
run(p, "  ★  RECOMMENDED FOR HAND-OVER  ★  ",
    bold=True, size=12, color=WHITE)

banner(doc, "APPROACH 3  ·  Hybrid — Mumbai-region, secure, scalable",
       fill=HEX_TEAL, size=16)
doc.add_paragraph()

sub_heading(doc,
    "Stack: Cloudflare Pages (Frontend) + Hostinger KVM 4 Mumbai (Backend) + Supabase Pro Mumbai (Database)",
    color=TEAL)

body(doc,
     "The balanced choice for going live at 1,000 concurrent. Frontend "
     "goes on Cloudflare's free CDN with unlimited bandwidth. Backend runs "
     "on a Mumbai VPS (KVM 4 — 4 vCPU / 16 GB RAM) which has comfortable "
     "headroom because the database load is fully offloaded. The "
     "PostgreSQL database lives on Supabase's managed service in AWS "
     "Mumbai — daily backups, 7-day point-in-time recovery, SOC 2 + ISO "
     "27001 compliance, and physical separation from the backend server. "
     "Costs 57% LESS than Approach 1 with the SAME managed database — "
     "saves Rs 1,95,000 over 3 years.")

info_box(doc,
    "WHY KVM 4 (NOT KVM 2) FOR 1,000 CONCURRENT",
    "Because the database is offloaded to Supabase, the VPS only runs "
    "the Node API + Nginx — but at 1,000 concurrent users, Node alone "
    "needs cluster mode across multiple cores to spread load. KVM 4 "
    "(4 vCPU / 16 GB RAM) gives PM2 4 worker processes and plenty of RAM "
    "headroom — comfortable for 1,000 concurrent on an I/O-bound API "
    "workload. KVM 2 (2 vCPU / 8 GB) would be tight at this scale. If "
    "load ever exceeds 1,000 concurrent, vertical upgrade to KVM 8 "
    "(~Rs 2,199/mo promo) is a one-click panel action.",
    label_color=GOLD, fill=HEX_GOLD_L, border="F59E0B")

heading(doc, "Where each piece lives", color=TEAL)
architecture_table(doc, [
    ("[FRONTEND]", HEX_FE, HEX_FE_L,
     "Cloudflare Pages",
     "React build on Cloudflare global CDN. Mumbai edge node serves Indian patients. Unlimited bandwidth.",
     "Rs 0 (free, unlimited)"),
    ("[BACKEND]", HEX_BE, HEX_BE_L,
     "Hostinger KVM 4 VPS (Mumbai)",
     "4 vCPU + 16 GB RAM + 200 GB NVMe. Runs Node API + Nginx + PM2 cluster (4 workers). Handles 1,000 concurrent.",
     "Rs 1,099/mo promo Y1-Y2, Rs 2,399/mo renewal"),
    ("[DATABASE]", HEX_DB, HEX_DB_L,
     "Supabase Pro Mumbai ($25/mo)",
     "Managed Postgres in AWS Mumbai. SOC 2 + ISO 27001. PgBouncer pooler scales to 1,000+ concurrent.",
     "$25/mo (~Rs 2,125)"),
    ("[DOMAIN]", "92400E", HEX_GOLD_L,
     "Hostinger registrar",
     "Free .com Y1 with VPS plan, then standard renewal.",
     "Rs 0 Y1, Rs 1,299 renewal"),
])

heading(doc, "3-Year cost breakdown", color=TEAL)
pricing_table(doc, [
    ("FE", HEX_FE, "Frontend — Cloudflare Pages (free, unlimited)",
     "Rs 0", "Rs 0", "Rs 0"),
    ("BE", HEX_BE,
     "Backend — Hostinger KVM 4 Mumbai (4 vCPU / 16 GB, 24-mo plan, incl. 18% GST)",
     "Rs 15,562", "Rs 15,562", "Rs 33,970"),
    ("DB", HEX_DB,
     "Database — Supabase Pro Mumbai $25/mo (SOC 2 + ISO 27001)",
     "Rs 25,500", "Rs 25,500", "Rs 25,500"),
    ("DOMAIN", "92400E",
     "Domain — .com (free Y1 with VPS, then renewal)",
     "Rs 0", "Rs 1,299", "Rs 1,299"),
    ("SEC", "DC2626",
     "SSL — Cloudflare free + Let's Encrypt on VPS",
     "Rs 0", "Rs 0", "Rs 0"),
    ("EXTRA", "64748B",
     "Off-site DB backup — Backblaze B2 ~10 GB",
     "Rs 612", "Rs 612", "Rs 612"),
    ("EXTRA", "64748B",
     "Email — Resend free tier (3,000/mo free)",
     "Rs 0", "Rs 0", "Rs 0"),
], ["TOTAL per year", "Rs 41,674", "Rs 42,973", "Rs 61,381"])

info_box(doc,
    "★  3-YEAR TOTAL  ·  Rs 1,46,028  (~Rs 4,056/month average)",
    "Per-patient cost at 1,000-2,000 users/day: ~Rs 0.20 - 0.40. The "
    "managed database in Mumbai is the same one Approach 1 uses — but "
    "this approach swaps Render Pro ($85/mo) for a Mumbai KVM 4 VPS at "
    "~1/6th the cost. Saves Rs 1,95,000 over 3 years with NO security "
    "trade-off compared to Approach 1, and proper headroom for 1,000 "
    "concurrent peak.",
    label_color=GOLD, fill=HEX_GOLD_L, border="F59E0B")

heading(doc, "Pros & Cons", color=TEAL)
pros_cons(doc, pros=[
    "Comfortable headroom for 1,000 concurrent — Node cluster across 4 cores + Supabase pgBouncer.",
    "Database isolated from Backend — a VPS compromise does NOT auto-leak patient data.",
    "Supabase Mumbai region keeps patient data in India (DPDP Act expectation).",
    "Supabase SOC 2 Type II + ISO 27001 — compliance certs we can show auditors.",
    "Cloudflare Frontend = unlimited bandwidth, free DDoS + WAF protection.",
    "Backend on Mumbai VPS — low-latency for Indian users (~5-15 ms to Supabase).",
    "Saves Rs 1,95,000 vs Approach 1 over 3 years with the same managed database.",
    "Vertically scales: KVM 4 -> KVM 8 (~Rs 2,199/mo) if backend load doubles.",
], cons=[
    "Two billing vendors (Hostinger INR + Supabase USD) — small FX risk on DB bill.",
    "Backend VPS still needs OS-level admin (patching, PM2 monitoring, SSH hardening).",
    "Network hop from VPS to Supabase adds ~5-15 ms latency (still well below 50 ms target).",
    "KVM 4 renewal price more than doubles in Year 3 (Rs 1,099 -> Rs 2,399/mo).",
    "Supabase Pro includes 250 GB egress per month — very heavy reports traffic could overflow (Rs 800/100 GB after).",
])

doc.add_page_break()


# ===================================================================
# FRONTEND DEEP-DIVE — Cloudflare Pages stays free forever
# ===================================================================
banner(doc, "FRONTEND DEEP-DIVE  ·  Stays Rs 0/month forever",
       fill=HEX_FE, size=15)
doc.add_paragraph()

sub_heading(doc,
    "Why the frontend never appears in the monthly bill (and never will, even at 10x scale)",
    color=TEAL)

body(doc,
     "The React frontend is just static files — HTML, JS, CSS bundles "
     "generated by `pnpm build`. Static files are served from a Content "
     "Delivery Network (CDN) — a global edge cache that scales horizontally "
     "to whatever traffic shows up. On Cloudflare Pages, this is free "
     "forever for the kind of traffic Arogya will generate. Below is "
     "exactly why, and how it compares to the current Netlify Free setup.")

info_box(doc,
    "✓  CLOUDFLARE PAGES FREE — VERIFIED LIMITS (May 2026)",
    "Bandwidth: UNLIMITED · Requests: UNLIMITED · Builds: 500/month (we'll "
    "use ~10-30) · Files: up to 20,000 (our React bundle has ~50-100) · "
    "Custom domains: 100 · Free SSL · Free unmetered DDoS protection "
    "(L3/L4/L7) · Free WAF · Global CDN with Mumbai edge POP. There is no "
    "credit card required, no bandwidth meter, no overage charges, no "
    "request limits. Cloudflare's only restriction is a fair-use clause "
    "for extremely high volumes (millions of requests per second) which "
    "Arogya will not hit at 1,000-2,000 daily users.",
    label_color=ORANGE, fill=HEX_ORANGE_L, border="F97316")

heading(doc, "Frontend bandwidth math — why we don't need to worry",
        color=TEAL)
tbl = doc.add_table(rows=6, cols=3)
hdrs = ["Metric", "Estimated value", "Notes"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
fe_math_rows = [
    ("React bundle size (gzipped)",
     "~800 KB - 1.2 MB",
     "Vite build, code-split, gzip+brotli compressed"),
    ("First load per visitor",
     "~1 MB (one-time)",
     "After that, browser cache + service worker serve repeat visits"),
    ("Daily unique visitors (peak)",
     "2,000",
     "Going-live scale upper bound"),
    ("Worst-case monthly bandwidth",
     "~60-90 GB",
     "If EVERY visitor was new each day for 30 days — impossible in reality"),
    ("Realistic monthly bandwidth",
     "~20-40 GB",
     "Most visitors return; bundle cached locally"),
]
for r_idx, row in enumerate(fe_math_rows, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10, color=SLATE,
            bold=(c_idx == 0))
        if r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

info_box(doc,
    "✓  EVEN AT 10x SCALE, FRONTEND IS STILL FREE",
    "At 20,000 daily users (10x of going-live target), realistic monthly "
    "frontend bandwidth would be ~200-400 GB. Cloudflare Pages Free still "
    "covers this — there is no bandwidth meter to overflow. Approach 3 "
    "could comfortably reach 10x scale before the frontend ever costs a "
    "rupee. (If Cloudflare ever flags fair-use, the Pages Pro plan starts "
    "at $20/mo — but Arogya will not need it.)",
    label_color=GREEN, fill=HEX_GREEN_L, border="16A34A")

heading(doc, "Netlify Free vs Cloudflare Pages Free — why we migrate",
        color=TEAL)
tbl = doc.add_table(rows=8, cols=3)
hdrs = ["Capability", "Netlify Free (current)", "Cloudflare Pages Free (target)"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=10, color=WHITE)
fe_compare_rows = [
    ("Monthly bandwidth",
     "100 GB hard cap",
     "UNLIMITED (no cap)"),
    ("Build minutes",
     "300/month",
     "500 builds/month (with 20-min timeout)"),
    ("Site-paused-on-overage risk",
     "YES — site pauses if 100 GB hit",
     "No — fair-use, no auto-pause"),
    ("DDoS protection",
     "Basic",
     "Unmetered L3/L4/L7 + free WAF"),
    ("Mumbai edge POP",
     "Yes",
     "Yes (Cloudflare has multiple India POPs)"),
    ("Custom domain SSL",
     "Free (Let's Encrypt)",
     "Free (Cloudflare SSL)"),
    ("Cost",
     "Rs 0",
     "Rs 0"),
]
for r_idx, row in enumerate(fe_compare_rows, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10, color=SLATE,
            bold=(c_idx == 0))
        if c_idx == 2:
            shade(cell, HEX_GOLD_L)
        elif r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

heading(doc, "Frontend migration plan (Netlify → Cloudflare Pages)",
        color=TEAL)
body(doc,
     "Cutover is ~15 minutes of work and ZERO downtime. The two CDNs can "
     "run in parallel while DNS switches over.")

migration_steps = [
    ("1. Sign in to Cloudflare dashboard",
     "dash.cloudflare.com → Workers & Pages → Create → Pages → Connect to Git"),
    ("2. Authorize the GitHub repo",
     "neurozenith26-create/Arogya_Clinic → Cloudflare reads commits + auto-deploys"),
    ("3. Configure build",
     "Build command: `pnpm --filter @arogya/frontend run build` · Output dir: `apps/frontend/dist` · Node 20"),
    ("4. Set env var",
     "VITE_API_URL = https://api.arogyadiagnostics.com (or current arogya-clinic.onrender.com)"),
    ("5. First deploy",
     "Cloudflare builds the React app and publishes to <project>.pages.dev — test it"),
    ("6. Add custom domain",
     "Cloudflare → Pages → Domains → Add arogyadiagnostics.com → DNS auto-configured"),
    ("7. Update backend CORS",
     "Set CORS_ORIGIN on the backend to the new Cloudflare URL (and keep Netlify URL until DNS fully propagates)"),
    ("8. Switch DNS",
     "Point arogyadiagnostics.com A/CNAME records to Cloudflare. TTL = 5 min for quick rollback if needed."),
    ("9. Keep Netlify running 24-48 hrs",
     "As DNS propagates globally, some users may still hit Netlify briefly — both work fine, no downtime."),
    ("10. Decommission Netlify",
     "After 48 hrs, delete the Netlify site to avoid confusion. Save the netlify.toml + screenshots in /docs."),
]
for s, d in migration_steps:
    tbl = doc.add_table(rows=1, cols=2)
    l = tbl.cell(0, 0)
    r = tbl.cell(0, 1)
    l.width = Cm(4.5)
    r.width = Cm(11.5)
    shade(l, HEX_FE)
    borders(l, color="0EA5E9", size="8")
    borders(r, color="E2E8F0", size="6")
    l.text = ""
    run(l.paragraphs[0], s, bold=True, size=10, color=WHITE)
    r.text = ""
    run(r.paragraphs[0], d, size=10, color=SLATE)
    doc.add_paragraph()

info_box(doc,
    "WHAT IF FRONTEND EVER NEEDS A PAID TIER?",
    "Cloudflare Pages Pro is $20/mo (Rs ~1,700/mo) and adds 5,000 builds/"
    "month + priority support — Arogya will NEVER need this at clinic "
    "scale. The only realistic upgrade path is if the React bundle starts "
    "calling Cloudflare Workers (server-side compute) heavily — but our "
    "backend is on Hostinger VPS, so Workers are not in our architecture. "
    "Frontend cost line stays at Rs 0/month for the lifetime of the app.",
    label_color=TEAL, fill=HEX_TEAL_L, border="0F766E")

doc.add_page_break()


# ===================================================================
# SIDE-BY-SIDE DETAILED COMPARISON
# ===================================================================
banner(doc, "SIDE-BY-SIDE COMPARISON",
       fill=HEX_ORANGE, size=16)
doc.add_paragraph()

heading(doc, "3.  3-Year cost summary (with Current Stack baseline)", color=TEAL)
tbl = doc.add_table(rows=5, cols=5)
hdrs = ["Approach", "Year 1", "Year 2", "Year 3", "3-Year Total"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
comp_rows = [
    ("0 · Current Stack (Netlify Free + Render Free + Supabase Free Mumbai)  ⚠ crashes live",
     "Rs 0", "Rs 0", "Rs 0", "Rs 0"),
    ("1 · Premium All-Managed (Cloudflare + Render Pro + Supabase Pro)",
     "Rs 1,13,011", "Rs 1,14,111", "Rs 1,14,111", "Rs 3,41,233"),
    ("2 · Single Hostinger KVM 4 VPS (all-in-one, tight at 1,000 concurrent)",
     "Rs 16,174", "Rs 17,473", "Rs 35,877", "Rs 69,524"),
    ("3 · Hybrid (Cloudflare + Hostinger KVM 4 + Supabase Pro) ★",
     "Rs 41,674", "Rs 42,973", "Rs 61,381", "Rs 1,46,028"),
]
for r_idx, row in enumerate(comp_rows, start=1):
    is_rec = "★" in row[0]
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        if is_rec:
            shade(cell, HEX_GOLD_L)
            run(cell.paragraphs[0], val, bold=True, size=10, color=TEAL_DARK)
        else:
            run(cell.paragraphs[0], val, size=10, color=SLATE,
                bold=(c_idx == 0))
            if r_idx % 2 == 1:
                shade(cell, HEX_GRAY_L)
doc.add_paragraph()

heading(doc, "4.  Decision matrix — who wins on what",
        color=TEAL)
tbl = doc.add_table(rows=10, cols=5)
hdrs = ["Criterion", "Approach 0", "Approach 1",
        "Approach 2", "Approach 3 ★"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=10, color=WHITE)

matrix = [
    ("Lowest 3-year cost",
     "★ (but unviable)", "✗ Costly", "★ Cheapest viable", "Balanced"),
    ("Handles 1,000 concurrent peak",
     "✗ Will crash", "★ Yes (Pro)", "Tight (one box)", "★ Yes"),
    ("Security for patient data (PHI/PII)",
     "Weak", "★ Winner", "Risky (one box)", "★ Winner"),
    ("Data residency in India",
     "✗ US/EU only", "DB only (Mumbai)", "Yes (Mumbai)", "Yes (Mumbai)"),
    ("Managed daily backups (7-day)",
     "✗ No (Supabase Free)", "★ Yes", "Manual cron", "★ Yes"),
    ("Compliance (SOC 2 / ISO 27001)",
     "✗ None", "★ Yes", "None", "★ Yes"),
    ("Easiest to operate (no devops)",
     "Medium", "★ Winner", "Needs Linux admin", "Middle"),
    ("Scalability to 2x-10x growth",
     "✗ Free tiers cap", "★ Yes", "Vertical only", "★ Yes"),
    ("Overall winner for Arogya",
     "Dev-only", "Too costly", "Risky for PHI", "★ Winner"),
]
for r_idx, row in enumerate(matrix, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        is_winner = "★" in val
        is_red = "✗" in val
        run(cell.paragraphs[0], val, size=9,
            color=GREEN if is_winner else (RED if is_red else SLATE),
            bold=(c_idx == 0 or is_winner))
        if is_winner:
            shade(cell, HEX_GREEN_L)
        elif is_red:
            shade(cell, HEX_RED_L)
        elif r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

doc.add_page_break()


# ===================================================================
# SECURITY COMPARISON
# ===================================================================
banner(doc, "SECURITY  ·  Patient Data Protection (DPDP Act 2023)",
       fill=HEX_TEAL, size=16)
doc.add_paragraph()

body(doc,
     "India's DPDP Act 2023 makes the clinic legally responsible for "
     "protecting patient PII and PHI. Below is how each approach handles "
     "the controls the law expects.")

tbl = doc.add_table(rows=11, cols=5)
hdrs = ["Security control", "Approach 0", "Approach 1", "Approach 2", "Approach 3 ★"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=10, color=WHITE)

sec_rows = [
    ("Encryption at rest (AES-256)",
     "Yes (Supabase)",
     "Yes (Supabase)",
     "Manual — must enable",
     "Yes (Supabase)"),
    ("Encryption in transit (TLS 1.2+)",
     "Yes (forced)",
     "Yes (forced)",
     "Manual (Let's Encrypt)",
     "Yes (Cloudflare + Supabase)"),
    ("DDoS protection",
     "Platform default only",
     "Built-in (Render edge)",
     "None — VPS IP exposed",
     "Cloudflare unmetered (free)"),
    ("Web Application Firewall (WAF)",
     "None",
     "Partial",
     "None unless installed",
     "Cloudflare free + 5 rules"),
    ("Database isolated from Backend",
     "Yes (separate service)",
     "Yes (separate service)",
     "No — same box",
     "Yes (Supabase vs VPS)"),
    ("Daily backups (7-day retention)",
     "None (Free tier)",
     "Yes (Supabase Pro)",
     "Manual pg_dump cron",
     "Yes (Supabase Pro)"),
    ("Data residency in India",
     "DB Mumbai (backend US)",
     "DB Mumbai (backend US)",
     "Yes (Mumbai)",
     "Yes (Mumbai both)"),
    ("Compliance certificates",
     "None (Free tier)",
     "SOC 2 + ISO 27001",
     "None",
     "SOC 2 + ISO 27001"),
    ("OS / DB patching",
     "Done for you",
     "Done for you",
     "Owner's responsibility",
     "DB auto; VPS only"),
    ("Audit logs",
     "Limited",
     "Yes (Supabase + Render)",
     "Manual setup (pgAudit)",
     "Yes (Supabase)"),
]
for r_idx, row in enumerate(sec_rows, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=9, color=SLATE,
            bold=(c_idx == 0))
        if c_idx == 4:
            shade(cell, HEX_GOLD_L)
        elif r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

info_box(doc,
    "ON HIPAA  ·  Honest note",
    "HIPAA is a US law and is NOT legally required for an Indian clinic "
    "serving Indian patients. India's DPDP Act 2023 is what applies, and "
    "the Approach 3 controls above satisfy it. If Arogya later expands to "
    "US patients/insurance, the database tier can be upgraded to a "
    "HIPAA-BAA plan as a one-time migration.",
    label_color=ORANGE, fill=HEX_ORANGE_L, border="F97316")

doc.add_page_break()


# ===================================================================
# CI/CD + RECOVERY (applies to all approaches)
# ===================================================================
banner(doc, "CI/CD & RECOVERY — How users stay online during code updates",
       fill=HEX_TEAL, size=15)
doc.add_paragraph()

heading(doc, "5.  Zero-downtime deploys (GitHub -> Live)", color=TEAL)
body(doc,
     "Code improvements and schema migrations go live without interrupting "
     "patients. Every approach uses the same pipeline:")

steps = [
    ("1. Push to GitHub main",
     "git push triggers GitHub Actions automatically."),
    ("2. CI runs tests",
     "lint + typecheck + unit tests + production build. Failure BLOCKS the deploy."),
    ("3. DB migrations applied",
     "Workflow runs `pnpm db:migrate` — only new migrations execute (tracked in _migrations table)."),
    ("4. Frontend deployed",
     "Static React build pushed to Cloudflare CDN globally in seconds."),
    ("5. Backend zero-downtime reload",
     "pm2 reload keeps OLD Node process serving in-flight requests until NEW one is ready, then swaps."),
    ("6. Smoke test + auto-rollback",
     "Workflow curls /api/v1/health. If non-200, reverts to previous build automatically."),
]
for s, d in steps:
    tbl = doc.add_table(rows=1, cols=2)
    l = tbl.cell(0, 0)
    r = tbl.cell(0, 1)
    l.width = Cm(4.5)
    r.width = Cm(11.5)
    shade(l, HEX_TEAL)
    borders(l, color="0F766E", size="8")
    borders(r, color="E2E8F0", size="6")
    l.text = ""
    run(l.paragraphs[0], s, bold=True, size=10, color=WHITE)
    r.text = ""
    run(r.paragraphs[0], d, size=10, color=SLATE)
    doc.add_paragraph()

info_box(doc,
    "USER IMPACT DURING DEPLOY  ·  0 seconds",
    "Patients booking tests or downloading reports do NOT see deploys. "
    "PM2 reload keeps the old version serving traffic until the new "
    "version is healthy. In-flight HTTP requests finish on the old "
    "process; new ones go to the new process. This applies to all 3 "
    "approaches.",
    label_color=TEAL, fill=HEX_TEAL_L, border="0F766E")

heading(doc, "6.  If a server goes down — recovery plan", color=TEAL)

dr_rows = [
    ("Node Backend crashes (bug/OOM)",
     "API errors ~5 sec",
     "PM2 auto-restart",
     "5-10 sec auto"),
    ("VPS / Render reboot",
     "Site unreachable briefly",
     "systemd / Render auto-restart",
     "1-2 min auto"),
    ("Database outage (managed)",
     "DB connections fail",
     "Supabase failover handles internally",
     "Minutes to hours"),
    ("Disk corruption / VPS dies",
     "Backend down (data safe on Supabase in App 1/3)",
     "Restore from snapshot or rebuild",
     "15-60 min manual"),
    ("Full disaster",
     "Everything down",
     "Re-provision + restore from PITR + Backblaze",
     "1-2 hours manual"),
]
tbl = doc.add_table(rows=len(dr_rows) + 1, cols=4)
hdrs = ["Failure", "What users see", "Recovery", "Time"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=10, color=WHITE)
for r_idx, row in enumerate(dr_rows, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=9, color=SLATE,
            bold=(c_idx == 0))
        if r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

# Uptime table
heading(doc, "7.  Realistic uptime expectation per approach", color=TEAL)
tbl = doc.add_table(rows=5, cols=3)
hdrs = ["Approach", "Uptime", "Downtime per month"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
upt = [
    ("0 · Current Stack (Free tier)", "~95% – 98%", "Cold starts + auto-pause risk"),
    ("1 · Premium All-Managed", "99.95%", "~22 minutes"),
    ("2 · Single Hostinger VPS", "99.5% – 99.9%", "~45 min - 3.5 hours"),
    ("3 · Hybrid ★", "99.7% – 99.95%", "~22 min - 2 hours"),
]
for r_idx, row in enumerate(upt, start=1):
    is_rec = "★" in row[0]
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10, color=TEAL_DARK if is_rec else SLATE,
            bold=(c_idx == 0))
        if is_rec:
            shade(cell, HEX_GOLD_L)
        elif r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

doc.add_page_break()


# ===================================================================
# FINAL RECOMMENDATION
# ===================================================================
banner(doc, "FINAL RECOMMENDATION",
       fill=HEX_GOLD, size=17)
doc.add_paragraph()

info_box(doc,
    "★  GO WITH APPROACH 3 — HYBRID (1,000-CONCURRENT, LIVE PRODUCTION)  ★",
    "For going live at 1,000-2,000 patients/day with up to 1,000 concurrent "
    "peak — no sleep, 3+ years — Approach 3 is the right balance: [FRONTEND] "
    "Cloudflare Pages (free, unlimited), [BACKEND] Hostinger KVM 4 Mumbai "
    "(4 vCPU / 16 GB RAM, Rs 1,099/mo promo), [DATABASE] Supabase Pro Mumbai "
    "($25/mo, SOC 2 + ISO 27001). 3-year cost Rs 1,46,028 — saves Rs 1,95,205 "
    "vs Approach 1 with the SAME managed database and the same 1,000-concurrent "
    "capacity, and adds proper database isolation that Approach 2 lacks.",
    label_color=GOLD, fill=HEX_GOLD_L, border="F59E0B")

info_box(doc,
    "✓  BUDGET FIT  ·  Client budget Rs 5,000-10,000 / month",
    "Approach 3 average monthly cost: Rs 3,473/mo (Year 1-2 on promo) → "
    "Rs 5,115/mo (Year 3 onwards at renewal price). Both years sit "
    "comfortably WITHIN the Rs 5,000-10,000/month budget with headroom to "
    "spare. Approach 1 (Rs 9,479/mo) is at the very top of budget. "
    "Approach 2 (Rs 1,931/mo) is well below budget but co-locates DB on "
    "the same box (not safe for patient data at this scale).",
    label_color=GREEN, fill=HEX_GREEN_L, border="16A34A")

heading(doc, "8a.  Monthly cost breakdown — Approach 3 (Recommended)",
        color=TEAL)
tbl = doc.add_table(rows=5, cols=3)
hdrs = ["Period", "Total Rs / month (avg)", "What's included"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
monthly_rows = [
    ("Year 1 (promo)",
     "~Rs 3,473 / month",
     "KVM 4 promo Rs 1,297 + Supabase $25 ≈ Rs 2,125 + Backblaze Rs 51 (domain free Y1)"),
    ("Year 2 (promo)",
     "~Rs 3,581 / month",
     "Same as Y1 + domain renewal Rs 108/mo amortized"),
    ("Year 3 (KVM renewal)",
     "~Rs 5,115 / month",
     "KVM 4 renewal Rs 2,831 + Supabase Rs 2,125 + domain Rs 108 + backups Rs 51"),
    ("Year 4 onwards (ongoing)",
     "~Rs 5,115 / month",
     "Same as Y3 — steady-state cost (no further price jumps)"),
]
for r_idx, row in enumerate(monthly_rows, start=1):
    is_highlight = r_idx == 3
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10,
            color=TEAL_DARK if is_highlight else SLATE,
            bold=(c_idx == 0 or c_idx == 1))
        if is_highlight:
            shade(cell, HEX_GOLD_L)
        elif r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

heading(doc, "8.  Rollout checklist", color=TEAL)
checklist = [
    "Register .com domain with Hostinger (free first year on 24-month VPS plan).",
    "Buy Hostinger KVM 4 24-month VPS plan in Mumbai DC at Rs 1,099/mo promo (4 vCPU, 16 GB RAM, 200 GB NVMe).",
    "Sign up for Supabase Pro in Mumbai (ap-south-1) region; run `pnpm db:migrate` on the 21 existing migrations.",
    "Connect GitHub repo to Cloudflare Pages — Frontend deploys automatically on push.",
    "Point domain DNS through Cloudflare for free DDoS + WAF + SSL.",
    "Harden the VPS: UFW firewall, fail2ban, SSH key-only auth, automatic security updates.",
    "Install Node 20, PM2 (cluster mode using all 4 vCPUs), Nginx reverse proxy, Let's Encrypt SSL.",
    "Set up GitHub Actions CI/CD: push to main -> tests -> SSH deploy with `pm2 reload` (zero-downtime).",
    "Configure nightly encrypted pg_dump to Backblaze B2 (~Rs 50/mo for 10 GB).",
    "Set up UptimeRobot (free) to ping the health endpoint every 5 minutes + email alerts.",
    "Plan a Year-2 migration of report PDFs from Postgres BYTEA -> Cloudflare R2 (free egress, $0.015/GB/mo) before the 8 GB DB ceiling is hit.",
    "Document the runbook in /docs/runbook.md and hand over credentials in Bitwarden/1Password.",
]
for i, step in enumerate(checklist, start=1):
    bullet(doc, f"{i}.  {step}", marker_color=GOLD, marker="✓")

doc.add_paragraph()

heading(doc, "9.  When to upgrade later (growth triggers)", color=TEAL)
tbl = doc.add_table(rows=6, cols=3)
hdrs = ["When this happens", "Do this", "New monthly cost"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
triggers = [
    ("Backend CPU > 70% sustained at peak",
     "Upgrade KVM 4 -> KVM 8 (8 vCPU, 32 GB) — one-click panel action.",
     "~Rs 2,199/mo (Y1 promo), Rs 4,399 renewal"),
    ("DB size approaches 8 GB (Supabase Pro ceiling)",
     "Migrate report PDFs from Postgres BYTEA to Cloudflare R2 object storage (FREE egress, $0.015/GB/mo).",
     "~Rs 1,500-5,000/yr for 100-300 GB"),
    ("DB connection limit hit (>200)",
     "Already using Supabase pooler; if still hit, upgrade to Team plan.",
     "$599/mo (Supabase Team)"),
    ("Reports egress > 250 GB/month",
     "Supabase Pro includes 250 GB; if past, pay Rs 800/100 GB extra (or move to R2 — free egress).",
     "+ Rs 800/100 GB extra"),
    ("Need 99.99% uptime SLA",
     "Add second VPS + HAProxy load balancer + Postgres replica.",
     "~Rs 2,500/mo extra"),
]
for r_idx, row in enumerate(triggers, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10, color=SLATE,
            bold=(c_idx == 0))
        if r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

doc.add_page_break()


# ===================================================================
# MAINTAINING THE LIVE APP — Development, Versions, Hotfixes
# ===================================================================
banner(doc, "MAINTAINING THE LIVE APP  ·  Development · Versions · Hotfixes",
       fill=HEX_TEAL, size=15)
doc.add_paragraph()

body(doc,
     "Once Approach 3 is live, this section explains how new features are "
     "developed, how versions are tracked, how the database schema evolves, "
     "and how urgent bug fixes get to production WITHOUT taking the site "
     "down. The whole workflow is built around three Git branches, two "
     "live environments, and one rule: production is sacred.")

heading(doc, "10.  Two-environment setup (Production + Staging)",
        color=TEAL)
body(doc,
     "Approach 3 runs TWO copies of the app — a live production stack the "
     "patients use, and a smaller staging stack the developer uses to "
     "rehearse every change before it touches production. Staging is "
     "identical in shape but uses cheaper tiers because no real patients "
     "see it.")

tbl = doc.add_table(rows=6, cols=3)
hdrs = ["Layer", "PRODUCTION (live)", "STAGING (rehearsal)"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
env_rows = [
    ("Domain",
     "arogyadiagnostics.com",
     "staging.arogyadiagnostics.com"),
    ("Frontend",
     "Cloudflare Pages (main branch)",
     "Cloudflare Pages (staging branch, free preview)"),
    ("Backend",
     "Hostinger KVM 4 Mumbai (port 3001)",
     "Same KVM 4 (port 3002) OR Hostinger KVM 1 (Rs 599/mo)"),
    ("Database",
     "Supabase Pro Mumbai ($25/mo) — REAL patient data",
     "Supabase Free Mumbai (separate project) — fake test data"),
    ("Monthly cost",
     "~Rs 5,115/mo (Year 3 steady-state)",
     "Rs 0 - 700/mo (Free Supabase + optional small VPS)"),
]
for r_idx, row in enumerate(env_rows, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10, color=SLATE,
            bold=(c_idx == 0))
        if c_idx == 1:
            shade(cell, HEX_GOLD_L)
        elif r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

info_box(doc,
    "GOLDEN RULE  ·  Real patient data NEVER leaves production",
    "Staging gets a one-time seeded copy of anonymized test data — never "
    "a dump of real patient records. A small `seed_staging.sql` script "
    "creates ~50 fake users, doctors, bookings, and reports. This protects "
    "PII/PHI even from the developer's own laptop, and satisfies DPDP "
    "Act expectations for data minimization.",
    label_color=RED, fill=HEX_RED_L, border="DC2626")


heading(doc, "11.  Branch strategy — how code flows from laptop to live",
        color=TEAL)
body(doc,
     "Three long-lived branches; every change merges UP the chain. "
     "Production never receives a direct push.")

steps = [
    ("feature/<name>",
     "Developer creates a branch off `development` for each new feature or fix. Work happens here, never directly on main."),
    ("development",
     "Default branch on GitHub. Feature branches merge here via PR. Auto-deploys to local dev environment when developer pulls."),
    ("staging",
     "Cherry-pick or merge from `development` when ready to rehearse against production-like infra. Auto-deploys to staging.arogyadiagnostics.com via GitHub Actions."),
    ("main",
     "PROTECTED branch. Only fast-forward merges from `staging` after staging has been smoke-tested. Auto-deploys to arogyadiagnostics.com via GitHub Actions."),
    ("hotfix/<bug>",
     "Branched directly from `main` for urgent production bugs. Cherry-picked back into `development` after merge."),
]
for s, d in steps:
    tbl = doc.add_table(rows=1, cols=2)
    l = tbl.cell(0, 0)
    r = tbl.cell(0, 1)
    l.width = Cm(4.5)
    r.width = Cm(11.5)
    shade(l, HEX_BE)
    borders(l, color="8B5CF6", size="8")
    borders(r, color="E2E8F0", size="6")
    l.text = ""
    run(l.paragraphs[0], s, bold=True, size=10, color=WHITE)
    r.text = ""
    run(r.paragraphs[0], d, size=10, color=SLATE)
    doc.add_paragraph()

info_box(doc,
    "BRANCH PROTECTION (GitHub setting)",
    "Set on the `main` branch: require pull request review, require status "
    "checks (CI tests must pass), require branches to be up to date, no "
    "direct pushes allowed, no force-pushes. This makes it physically "
    "impossible to deploy untested code to production by accident.",
    label_color=TEAL, fill=HEX_TEAL_L, border="0F766E")


heading(doc, "12.  Database migrations & schema versioning", color=TEAL)
body(doc,
     "The database schema is versioned in SQL files under "
     "`supabase/migrations/0001_*.sql` ... `0021_*.sql` (21 existing). "
     "Every schema change is a NEW migration file — never edit an old one. "
     "Applied migrations are tracked in a `_migrations` table so each one "
     "runs exactly once.")

mig_rows = [
    ("1. Write the migration locally",
     "Create `supabase/migrations/0022_add_column.sql` with the forward-only DDL (ALTER TABLE, CREATE INDEX, etc.)"),
    ("2. Test against local Postgres",
     "Spin up a local Postgres in Docker, run `pnpm db:migrate`. Verify the schema looks right and the app still works."),
    ("3. Apply to staging",
     "Merge to `staging` branch → GitHub Actions runs `pnpm db:migrate` against staging Supabase. Smoke test."),
    ("4. Apply to production",
     "Merge to `main` → CI runs `pnpm db:migrate` against production Supabase BEFORE the new backend code is deployed."),
    ("5. Backend code deploy",
     "PM2 reload picks up the new code with the new schema already in place. Zero downtime."),
    ("6. If migration fails",
     "CI auto-stops — the old backend code keeps running against the old schema. Developer rolls forward with a fix migration (never rolls back the DB)."),
]
for s, d in mig_rows:
    tbl = doc.add_table(rows=1, cols=2)
    l = tbl.cell(0, 0)
    r = tbl.cell(0, 1)
    l.width = Cm(4.5)
    r.width = Cm(11.5)
    shade(l, HEX_DB)
    borders(l, color="059669", size="8")
    borders(r, color="E2E8F0", size="6")
    l.text = ""
    run(l.paragraphs[0], s, bold=True, size=10, color=WHITE)
    r.text = ""
    run(r.paragraphs[0], d, size=10, color=SLATE)
    doc.add_paragraph()

info_box(doc,
    "FORWARD-ONLY MIGRATIONS",
    "Never roll back a database migration in production. If a migration "
    "breaks something, write a NEW migration that fixes it. This sounds "
    "harsh but it's the industry standard — rolling back can corrupt data "
    "if any rows were written under the new schema. Daily Supabase backups "
    "(7-day retention) cover the worst case.",
    label_color=ORANGE, fill=HEX_ORANGE_L, border="F97316")


heading(doc, "13.  Version tagging & changelog", color=TEAL)
body(doc,
     "Every production release gets a Git tag and an entry in CHANGELOG.md. "
     "This lets us answer 'what changed between Tuesday and today?' without "
     "reading commits one by one.")

ver_rows = [
    ("Release tags",
     "v1.0.0 (going live), v1.0.1 (bug fix), v1.1.0 (new feature), v2.0.0 (breaking change). Pattern: MAJOR.MINOR.PATCH (semver)."),
    ("Git tag command",
     "git tag -a v1.0.1 -m 'Fix UPI overcharge' && git push origin v1.0.1 — done at the moment a release is merged to main."),
    ("GitHub Releases",
     "Each tag becomes a GitHub Release with a copy of CHANGELOG.md notes attached. Auditable history forever."),
    ("Backend version endpoint",
     "GET /api/v1/version returns { tag: 'v1.0.1', sha: '67abd1d', built_at: '...' }. Frontend shows this in the footer."),
    ("Hand-over",
     "Any new developer can read CHANGELOG.md and reconstruct what shipped when, with no tribal knowledge required."),
]
tbl = doc.add_table(rows=len(ver_rows) + 1, cols=2)
hdrs = ["Practice", "How it works"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
for r_idx, row in enumerate(ver_rows, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10, color=SLATE,
            bold=(c_idx == 0))
        if r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()


heading(doc, "14.  Hotfix process — urgent production bug, no panic",
        color=TEAL)
body(doc,
     "A patient reports the booking flow is broken at 2 PM. Here is the "
     "exact sequence to get a fix live without disturbing other users.")

hotfix_rows = [
    ("T+0 min",
     "Triage",
     "Reproduce in staging if possible. Confirm the bug is in code (not data). Open a GitHub issue."),
    ("T+5 min",
     "Branch from main",
     "git checkout main && git pull && git checkout -b hotfix/booking-422"),
    ("T+5-30 min",
     "Write & test the fix",
     "Smallest possible change. Add a regression test that fails BEFORE the fix and passes AFTER."),
    ("T+30 min",
     "PR to main",
     "Open PR directly to main with the hotfix branch. CI runs full test suite + build."),
    ("T+35 min",
     "Auto-deploy",
     "Merge → GitHub Actions runs db:migrate (no-op if no DB change) → SSH deploy → pm2 reload → health check. Zero downtime."),
    ("T+40 min",
     "Tag & note",
     "git tag v1.0.2 -m 'hotfix: booking 422 error'. Add a CHANGELOG.md entry. Cherry-pick the fix back into `development` branch."),
    ("T+45 min",
     "Notify patient",
     "Confirm the fix in production. Reply to the original bug report."),
]
tbl = doc.add_table(rows=len(hotfix_rows) + 1, cols=3)
hdrs = ["Time", "Step", "What you do"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
for r_idx, row in enumerate(hotfix_rows, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10, color=SLATE,
            bold=(c_idx == 0 or c_idx == 1))
        if r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

info_box(doc,
    "ROLLBACK IF THE HOTFIX MAKES THINGS WORSE",
    "Every deploy keeps the previous Git SHA pinned on the VPS. To roll "
    "back instantly: `git checkout <previous-sha> && pm2 reload` — under "
    "30 seconds, zero data loss (DB schema unchanged because migrations "
    "are forward-only). For database-level rollback, restore from the "
    "most recent Supabase daily backup (within 7-day retention window).",
    label_color=GOLD, fill=HEX_GOLD_L, border="F59E0B")


heading(doc, "15.  Day-to-day maintenance cadence", color=TEAL)
maint_rows = [
    ("Daily (automatic)",
     "Supabase daily backup runs (no action needed) · UptimeRobot pings health endpoint every 5 min · pg_dump to Backblaze B2 nightly"),
    ("Weekly (15 min)",
     "Check UptimeRobot alerts inbox · Review Supabase usage dashboard · `pm2 logs` for any new errors"),
    ("Monthly (1 hour)",
     "Run `apt update && apt upgrade` on the VPS · Rotate Supabase database password · Review last 30 days of CHANGELOG · Pay invoices (Hostinger + Supabase)"),
    ("Quarterly (2-4 hours)",
     "Test the disaster-recovery runbook — restore staging from a Backblaze backup. Verify SSL certs auto-renewed. Audit IAM/SSH access."),
    ("Annually (1 day)",
     "Hostinger VPS plan renewal decision (KVM 4 → KVM 8 if growing). Domain renewal. Security audit. Update CLAUDE.md / runbook for any new conventions."),
]
tbl = doc.add_table(rows=len(maint_rows) + 1, cols=2)
hdrs = ["Cadence", "What happens"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=11, color=WHITE)
for r_idx, row in enumerate(maint_rows, start=1):
    for c_idx, val in enumerate(row):
        cell = tbl.rows[r_idx].cells[c_idx]
        cell.text = ""
        run(cell.paragraphs[0], val, size=10, color=SLATE,
            bold=(c_idx == 0))
        if r_idx % 2 == 1:
            shade(cell, HEX_GRAY_L)
doc.add_paragraph()

info_box(doc,
    "TOTAL MAINTENANCE TIME PER MONTH  ·  ~3-4 hours",
    "Most of the heavy lifting (backups, deploys, SSL renewal, security "
    "patches on Supabase / Cloudflare) is automated or provider-managed. "
    "The owner-developer's hands-on time is roughly 3-4 hours per month "
    "of routine work, plus whatever is spent building new features.",
    label_color=GREEN, fill=HEX_GREEN_L, border="16A34A")

doc.add_page_break()


# ===================================================================
# VERIFIED SOURCES
# ===================================================================
banner(doc, "VERIFIED SOURCES — every price confirmed May 2026",
       fill=HEX_TEAL, size=14)
doc.add_paragraph()

body(doc,
     "Every rupee and dollar figure in this document was verified from the "
     "provider's official website. The client can re-verify at any time.",
     italic=True)

sources = [
    ("Hostinger India VPS",
     "https://www.hostinger.com/in/vps-hosting",
     "VERIFIED May 2026: KVM 1 Rs 599 promo/Rs 999 renewal · KVM 2 Rs 799/Rs 1,199 · KVM 4 Rs 1,099/Rs 2,399 · KVM 8 Rs 2,199/Rs 4,399 (24-mo plan, 18% GST extra, AMD EPYC, NVMe SSD)"),
    ("Hostinger India domains",
     "https://www.hostinger.com/in/pricing/domains",
     "VERIFIED: .com Rs 199 promo Y1, Rs 1,299 renewal · FREE .com first year when bought with VPS plan"),
    ("Cloudflare Pages limits",
     "https://developers.cloudflare.com/pages/platform/limits",
     "Free tier: UNLIMITED bandwidth, 500 builds/month, 20K files max, 20-min build timeout"),
    ("Cloudflare Free DDoS + WAF",
     "https://www.cloudflare.com/plans/free",
     "Unmetered DDoS protection L3/L4/L7, basic WAF, free SSL, global CDN"),
    ("Supabase regions",
     "https://supabase.com/docs/guides/platform/regions",
     "Mumbai (ap-south-1) confirmed available · Pro plan supports all regions including Mumbai"),
    ("Supabase pricing",
     "https://supabase.com/pricing",
     "VERIFIED May 2026: Pro $25/mo · 8 GB DB · 250 GB egress · daily backups 7-day retention · ALWAYS-ON (never pauses) · PITR is $100/mo add-on (NOT included) · SOC 2 Type II"),
    ("Render pricing",
     "https://render.com/pricing",
     "VERIFIED May 2026: Starter $7/mo (0.5 CPU/512 MB) · Standard $25/mo (1 CPU/2 GB) · Pro $85/mo (2 CPU/4 GB) · Pro Plus $185/mo (4 CPU/8 GB) · Paid plans never sleep (only Free sleeps)"),
    ("Neon regions (for reference)",
     "https://neon.com/docs/introduction/regions",
     "Asia-Pacific: only Singapore + Sydney as of May 2026. Mumbai NOT available."),
    ("Aiven Developer Tier",
     "https://aiven.io/developer-tier",
     "$5/mo · 'For test and personal projects' only · NO region selection · NO PITR · NO HIPAA"),
    ("Backblaze B2 pricing",
     "https://www.backblaze.com/cloud-storage/pricing",
     "VERIFIED May 2026: $6.95/TB/mo (~$0.007/GB) · Free egress up to 3x avg storage · Unlimited free egress via Cloudflare CDN"),
    ("UptimeRobot",
     "https://uptimerobot.com",
     "Free tier: 50 monitors, 5-minute interval, email/SMS alerts"),
    ("Resend email",
     "https://resend.com/pricing",
     "Free tier: 3,000 emails/month, 100/day"),
    ("DPDP Act 2023 (India)",
     "https://www.meity.gov.in",
     "Ministry of Electronics and IT — Digital Personal Data Protection Act 2023"),
    ("DPDP Act in healthcare",
     "https://www.aarnalaw.com/insights/the-applicability-of-the-dpdp-act-in-hospitals",
     "Compliance guidance for clinics and hospitals"),
]
tbl = doc.add_table(rows=len(sources) + 1, cols=3)
hdrs = ["Source", "URL", "Verified fact"]
for i, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[i]
    shade(cell, HEX_TEAL)
    cell.text = ""
    run(cell.paragraphs[0], h, bold=True, size=10, color=WHITE)
for r_idx, (p, url, fact) in enumerate(sources, start=1):
    c0 = tbl.rows[r_idx].cells[0]
    c1 = tbl.rows[r_idx].cells[1]
    c2 = tbl.rows[r_idx].cells[2]
    c0.text = ""
    run(c0.paragraphs[0], p, bold=True, size=9, color=SLATE)
    c1.text = ""
    run(c1.paragraphs[0], url, size=8, color=FE)
    c2.text = ""
    run(c2.paragraphs[0], fact, size=8, color=SLATE)
    if r_idx % 2 == 1:
        shade(c0, HEX_GRAY_L)
        shade(c1, HEX_GRAY_L)
        shade(c2, HEX_GRAY_L)

doc.add_paragraph()
info_box(doc,
    "FX & TAX ASSUMPTIONS",
    "USD prices converted at USD 1 = INR 85 (May 2026 reference rate). "
    "Hostinger Indian prices include 18% GST. Supabase / Render / Backblaze "
    "are billed in USD by international providers — the clinic may need to "
    "handle GST on reverse-charge for proper tax compliance.",
    label_color=GRAY, fill=HEX_GRAY_L, border="64748B")


# ===================================================================
# SAVE  (into the project's Docs/ folder)
# ===================================================================
import os
docs_dir = "Docs"
os.makedirs(docs_dir, exist_ok=True)
output_path = os.path.join(docs_dir, "Arogya_Hosting_CLIENT_3Approaches.docx")
try:
    doc.save(output_path)
except PermissionError:
    output_path = os.path.join(docs_dir,
                               "Arogya_Hosting_CLIENT_3Approaches_v2.docx")
    doc.save(output_path)
    print("Primary was open in Word — wrote v2 instead.")
print(f"DONE: {os.path.abspath(output_path)}")
